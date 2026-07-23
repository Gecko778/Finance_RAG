"""文档解析：PDF / DOCX / TXT / MD → 纯文本。"""

from io import BytesIO
from pathlib import PurePosixPath

from docx import Document as DocxDocument
from pypdf import PdfReader


class UnsupportedFormatError(ValueError):
    pass


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(data: bytes) -> str:
    doc = DocxDocument(BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _parse_plain(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("gbk")


PARSERS = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".txt": _parse_plain,
    ".md": _parse_plain,
}


def parse_bytes(filename: str, data: bytes) -> str:
    suffix = PurePosixPath(filename).suffix.lower()
    parser = PARSERS.get(suffix)
    if parser is None:
        raise UnsupportedFormatError(f"不支持的格式：{suffix}（支持 {', '.join(PARSERS)}）")
    text = parser(data).strip()
    if not text:
        raise ValueError("解析结果为空（可能是扫描件或纯图片 PDF）")
    return text

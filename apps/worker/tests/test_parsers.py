"""解析器单元测试：TXT/MD/DOCX 用生成文件实测；PDF 走 M2 验收（真实政策文件）。"""

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from rag_worker.parsers import UnsupportedFormatError, parse_bytes


def test_parse_txt_utf8():
    assert parse_bytes("a.txt", "增值税政策".encode()) == "增值税政策"


def test_parse_txt_gbk_fallback():
    assert parse_bytes("a.txt", "增值税政策".encode("gbk")) == "增值税政策"


def test_parse_md():
    assert "第一条" in parse_bytes("a.md", "# 标题\n第一条 内容".encode())


def test_parse_docx_paragraphs_and_table():
    doc = DocxDocument()
    doc.add_paragraph("第一条 免征增值税。")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "税率"
    table.rows[0].cells[1].text = "3%"
    buf = BytesIO()
    doc.save(buf)

    text = parse_bytes("policy.docx", buf.getvalue())
    assert "第一条 免征增值税。" in text
    assert "税率\t3%" in text


def test_unsupported_format():
    with pytest.raises(UnsupportedFormatError):
        parse_bytes("a.xlsx", b"whatever")


def test_empty_result_raises():
    with pytest.raises(ValueError, match="解析结果为空"):
        parse_bytes("a.txt", b"   ")
